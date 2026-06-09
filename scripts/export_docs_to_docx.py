#!/usr/bin/env python3
"""One-time export: docs/*.md -> docs/docx/*.docx"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    print("Install dependency: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "docx"

EXPORTS = [
    ("README.md", "00_DOCUMENTATION_INDEX.docx"),
    ("01_SETUP_GUIDE.md", "01_SETUP_GUIDE.docx"),
    ("02_ARCHITECTURE.md", "02_ARCHITECTURE.docx"),
    ("03_API_DOCUMENTATION.md", "03_API_DOCUMENTATION.docx"),
    ("04_ENVIRONMENT_FLAVOR_GUIDE.md", "04_ENVIRONMENT_FLAVOR_GUIDE.docx"),
    ("05_FEATURE_DOCUMENTATION.md", "05_FEATURE_DOCUMENTATION.docx"),
    ("06_RELEASE_CHECKLIST.md", "06_RELEASE_CHECKLIST.docx"),
    ("07_PLAY_STORE_DEPLOYMENT.md", "07_PLAY_STORE_DEPLOYMENT.docx"),
    ("08_GIT_WORKFLOW_GUIDE.md", "08_GIT_WORKFLOW_GUIDE.docx"),
    ("09_TESTING_GUIDE.md", "09_TESTING_GUIDE.docx"),
]


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    i = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_rows = [parse_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(table_rows):
                for ci in range(cols):
                    val = row[ci] if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = val
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- [ ] "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("- [x] ") or line.startswith("- [X] "):
            doc.add_paragraph("☑ " + line[6:].strip(), style="List Bullet")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line).strip(), style="List Number")
        elif line.strip() == "---":
            doc.add_paragraph("─" * 40)
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            # Inline code `text`
            parts = re.split(r"(`[^`]+`)", line)
            for part in parts:
                if part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                else:
                    p.add_run(part)

        i += 1

    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Wrote {docx_path.relative_to(ROOT)}")


def main() -> None:
    for md_name, docx_name in EXPORTS:
        md_path = DOCS / md_name
        if not md_path.exists():
            print(f"Missing {md_path}")
            sys.exit(1)
        convert_md_to_docx(md_path, OUT / docx_name)
    print("Done.")


if __name__ == "__main__":
    main()
